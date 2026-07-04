import os
import subprocess
import logging
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from django.conf import settings

logger = logging.getLogger(__name__)

class GenerateDocumentService:
    def __init__(self):
        self.doc = None
        self.slate_900 = RGBColor(15, 23, 42)
        self.slate_700 = RGBColor(51, 65, 85)
        self.slate_600 = RGBColor(71, 85, 105)
        self.slate_500 = RGBColor(100, 116, 139)
        self.slate_400 = RGBColor(148, 163, 184)
        self.slate_200 = RGBColor(226, 232, 240)
        self.template = "modern"
        self.font_name = "Arial"

    def _set_paragraph_border(self, paragraph):
        p = paragraph._p
        pPr = p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '4')  # 1/2 pt
        bottom.set(qn('w:space'), '1')
        bottom.set(qn('w:color'), 'E2E2E2')
        pBdr.append(bottom)
        pPr.append(pBdr)

    def _add_section_header(self, text):
        h = self.doc.add_paragraph()
        h.paragraph_format.space_before = Pt(18 if self.template != 'minimal' else 12)
        h.paragraph_format.space_after = Pt(6 if self.template != 'minimal' else 4)
        
        if self.template == 'classic':
            h.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = h.add_run(text.upper())
            run.bold = True
            run.font.size = Pt(9.5)
            run.font.color.rgb = self.slate_900
            run.font.name = self.font_name
        elif self.template == 'minimal':
            run = h.add_run(text.upper())
            run.bold = True
            run.font.size = Pt(8.5)
            run.font.color.rgb = self.slate_900
            run.font.name = self.font_name
        else: # Modern
            self._set_paragraph_border(h)
            run = h.add_run(text.upper())
            run.bold = True
            run.font.size = Pt(9)
            run.font.color.rgb = self.slate_500
            run.font.name = self.font_name
            
            # Tracking/Kerning
            rPr = run._element.get_or_add_rPr()
            spacing = OxmlElement('w:spacing')
            spacing.set(qn('w:val'), '35')
            rPr.append(spacing)

    def generate(self, context=None):
        if context is None:
            context = DUMMY_RESUME_CONTEXT
            
        self.template = context.get('template', 'modern') or 'modern'
        self.font_name = 'Georgia' if self.template == 'classic' else 'Arial'
        
        self.doc = Document()
        
        # Set Default Style
        style = self.doc.styles['Normal']
        font = style.font
        font.name = self.font_name
        font.size = Pt(9.5 if self.template == 'minimal' else 10)
        font.color.rgb = self.slate_700

        # Set Margins (A4: 8.27 x 11.69 inches)
        sections = self.doc.sections
        for section in sections:
            margin = Inches(0.7 if self.template == 'minimal' else 1.0)
            section.top_margin = margin
            section.bottom_margin = margin
            section.left_margin = margin
            section.right_margin = margin

        # Name
        name_p = self.doc.add_paragraph()
        if self.template == 'classic':
            name_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
        name_text = context.get('full_name', '').upper() if self.template != 'minimal' else context.get('full_name', '')
        name_run = name_p.add_run(name_text)
        name_run.bold = True
        name_run.font.size = Pt(24 if self.template == 'classic' else (20 if self.template == 'minimal' else 28))
        name_run.font.name = self.font_name
        name_run.font.color.rgb = self.slate_900
        name_p.paragraph_format.space_after = Pt(4 if self.template == 'minimal' else 6)

        # Contact Info
        contact_items = []
        if context.get('email'): contact_items.append(context['email'])
        if context.get('phone_number'): contact_items.append(context['phone_number'])
        if context.get('location'): contact_items.append(context['location'])

        if self.template == 'classic':
            contact_p = self.doc.add_paragraph()
            contact_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            contact_text = "  •  ".join(contact_items)
            run = contact_p.add_run(contact_text)
            run.font.size = Pt(9.5)
            run.font.name = self.font_name
            run.font.color.rgb = self.slate_600
            contact_p.paragraph_format.space_after = Pt(16)
        elif self.template == 'minimal':
            contact_p = self.doc.add_paragraph()
            contact_text = "  |  ".join(contact_items)
            run = contact_p.add_run(contact_text)
            run.font.size = Pt(9)
            run.font.name = self.font_name
            run.font.color.rgb = self.slate_600
            contact_p.paragraph_format.space_after = Pt(12)
        else: # Modern
            for i, item in enumerate(contact_items):
                p = self.doc.add_paragraph()
                run = p.add_run(item)
                run.font.size = Pt(9.5)
                run.font.name = self.font_name
                run.font.color.rgb = self.slate_600
                
                # Use small spacing between items, but larger after the last one
                if i == len(contact_items) - 1:
                    p.paragraph_format.space_after = Pt(16)
                else:
                    p.paragraph_format.space_after = Pt(2)

        # Professional Summary
        if context.get('skill_description'):
            self._add_section_header("Professional Summary")
            summary_p = self.doc.add_paragraph()
            summary_p.paragraph_format.line_spacing = 1.2
            summary_p.paragraph_format.space_after = Pt(8 if self.template == 'minimal' else 12)
            summary_run = summary_p.add_run(context['skill_description'])
            summary_run.font.size = Pt(9.5 if self.template == 'minimal' else 10)
            summary_run.font.name = self.font_name

        # Experience
        if context.get('experiences'):
            self._add_section_header("Professional Experience")
            for exp in context['experiences']:
                # Combine Title and Company into ONE table with TWO rows
                table_width = Inches(6.87 if self.template == 'minimal' else 6.5)
                table = self.doc.add_table(rows=2, cols=2)
                table.width = table_width
                
                # Row 0: Job Title & Dates
                cells = table.rows[0].cells
                title_p = cells[0].paragraphs[0]
                title_p.paragraph_format.space_before = Pt(0)
                title_p.paragraph_format.space_after = Pt(0)
                title_run = title_p.add_run(exp.get('job_title', ''))
                title_run.bold = True
                title_run.font.size = Pt(10.5 if self.template == 'minimal' else 11.5)
                title_run.font.name = self.font_name
                title_run.font.color.rgb = self.slate_900
                
                date_p = cells[1].paragraphs[0]
                date_p.paragraph_format.space_before = Pt(0)
                date_p.paragraph_format.space_after = Pt(0)
                date_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                date_run = date_p.add_run(f"{exp.get('date_from', '')} — {exp.get('date_to') or 'Present'}")
                date_run.font.size = Pt(8 if self.template == 'minimal' else 8.5)
                date_run.font.name = self.font_name
                date_run.font.color.rgb = self.slate_500
                
                # Row 1: Company & Location
                cells2 = table.rows[1].cells
                comp_p = cells2[0].paragraphs[0]
                comp_p.paragraph_format.space_before = Pt(0)
                comp_p.paragraph_format.space_after = Pt(0)
                comp_run = comp_p.add_run(exp.get('company_name', ''))
                comp_run.bold = True
                comp_run.font.size = Pt(9.5 if self.template == 'minimal' else 10)
                comp_run.font.name = self.font_name
                comp_run.font.color.rgb = self.slate_700
                
                loc_p = cells2[1].paragraphs[0]
                loc_p.paragraph_format.space_before = Pt(0)
                loc_p.paragraph_format.space_after = Pt(0)
                loc_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                loc_run = loc_p.add_run(exp.get('location', ''))
                loc_run.italic = True
                loc_run.font.size = Pt(8 if self.template == 'minimal' else 8.5)
                loc_run.font.name = self.font_name
                loc_run.font.color.rgb = self.slate_400

                # Add a small margin at the bottom of the table
                spacer = self.doc.add_paragraph()
                spacer.paragraph_format.space_after = Pt(6 if self.template == 'minimal' else 8)
                spacer.paragraph_format.line_spacing = Pt(1)

                # Bullet Points
                for bp in exp.get('bullet_points', []):
                    if bp.strip():
                        b_p = self.doc.add_paragraph(style='List Bullet')
                        b_p.paragraph_format.left_indent = Inches(0.25)
                        b_p.paragraph_format.first_line_indent = Inches(-0.15)
                        b_p.paragraph_format.space_before = Pt(1 if self.template == 'minimal' else 2)
                        b_p.paragraph_format.space_after = Pt(1 if self.template == 'minimal' else 2)
                        
                        b_run = b_p.add_run(bp.strip())
                        b_run.font.size = Pt(9.5 if self.template == 'minimal' else 10)
                        b_run.font.name = self.font_name
                        b_run.font.color.rgb = self.slate_700
                
                self.doc.add_paragraph().paragraph_format.space_after = Pt(4 if self.template == 'minimal' else 6)

        # Education
        if context.get('educations'):
            self._add_section_header("Education")
            for edu in context['educations']:
                table_width = Inches(6.87 if self.template == 'minimal' else 6.5)
                table = self.doc.add_table(rows=2, cols=2)
                table.width = table_width
                
                # Row 0: School & Dates
                cells = table.rows[0].cells
                sch_p = cells[0].paragraphs[0]
                sch_p.paragraph_format.space_before = Pt(0)
                sch_p.paragraph_format.space_after = Pt(0)
                sch_run = sch_p.add_run(edu.get('school', ''))
                sch_run.bold = True
                sch_run.font.size = Pt(9.5 if self.template == 'minimal' else 10.5)
                sch_run.font.name = self.font_name
                sch_run.font.color.rgb = self.slate_900
                
                date_p = cells[1].paragraphs[0]
                date_p.paragraph_format.space_before = Pt(0)
                date_p.paragraph_format.space_after = Pt(0)
                date_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                date_run = date_p.add_run(f"{edu.get('date_from', '')} — {edu.get('date_to', '')}")
                date_run.font.size = Pt(8 if self.template == 'minimal' else 8.5)
                date_run.font.name = self.font_name
                date_run.font.color.rgb = self.slate_500
                
                # Row 1: School Type & Location
                cells2 = table.rows[1].cells
                type_p = cells2[0].paragraphs[0]
                type_p.paragraph_format.space_before = Pt(0)
                type_p.paragraph_format.space_after = Pt(0)
                type_run = type_p.add_run(edu.get('school_type', ''))
                type_run.font.size = Pt(9.5 if self.template == 'minimal' else 10)
                type_run.font.name = self.font_name
                
                loc_p = cells2[1].paragraphs[0]
                loc_p.paragraph_format.space_before = Pt(0)
                loc_p.paragraph_format.space_after = Pt(0)
                loc_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                loc_run = loc_p.add_run(edu.get('location', ''))
                loc_run.italic = True
                loc_run.font.size = Pt(8 if self.template == 'minimal' else 8.5)
                loc_run.font.name = self.font_name
                loc_run.font.color.rgb = self.slate_400
                
                self.doc.add_paragraph().paragraph_format.space_after = Pt(3 if self.template == 'minimal' else 4)

        # Skills
        if context.get('skills'):
            self._add_section_header("Skills")
            
            # Split skills into two columns
            skills = context.get('skills', [])
            mid = (len(skills) + 1) // 2
            left_col = skills[:mid]
            right_col = skills[mid:]
            
            # Create a two-column table for skills
            table = self.doc.add_table(rows=1, cols=2)
            table.autofit = False
            
            # Set column widths
            for column in table.columns:
                column.width = Inches(3.435 if self.template == 'minimal' else 3.25)
            
            def add_skills_to_cell(cell, skill_list):
                for i, skill in enumerate(skill_list):
                    if i == 0:
                        p = cell.paragraphs[0]
                    else:
                        p = cell.add_paragraph()
                    
                    p.style = 'List Bullet'
                    p.paragraph_format.left_indent = Inches(0.2)
                    p.paragraph_format.first_line_indent = Inches(-0.15)
                    p.paragraph_format.space_before = Pt(1)
                    p.paragraph_format.space_after = Pt(1)
                    
                    run = p.add_run(skill)
                    run.font.size = Pt(9.5 if self.template == 'minimal' else 10)
                    run.font.name = self.font_name
                    run.font.color.rgb = self.slate_700

            add_skills_to_cell(table.cell(0, 0), left_col)
            add_skills_to_cell(table.cell(0, 1), right_col)

        # Save & Output
        output_dir = os.path.join(settings.MEDIA_ROOT, 'resumes')
        if not os.path.exists(output_dir):
            os.makedirs(output_dir)
            
        full_name = context.get('full_name') or 'resume'
        full_name_slug = full_name.strip().lower().replace(' ', '_')
        output_path = os.path.join(output_dir, f'{full_name_slug}.docx')
        self.doc.save(output_path)
        
        logger.info(f"Generated DOCX: {output_path}")
        self.convert_to_pdf(output_path, output_dir)
        return full_name_slug
        
    def convert_to_pdf(self, docx_path, output_dir):
        try:
            cmd = ['libreoffice', '--headless', '--convert-to', 'pdf', docx_path, '--outdir', output_dir]
            if os.name == 'nt':
                try:
                    subprocess.run(cmd, check=True, capture_output=True)
                except FileNotFoundError:
                    cmd[0] = 'soffice'
                    subprocess.run(cmd, check=True, capture_output=True)
            else:
                subprocess.run(cmd, check=True, capture_output=True)
        except Exception as e:
            logger.error(f"PDF conversion failed: {e}")

DUMMY_RESUME_CONTEXT = {
    "full_name": "CARL JEFFERSON DELFIN",
    "email": "carljefferson.delfin@gmail.com",
    "phone_number": "09103594750",
    "location": "Rodriguez, Rizal, Philippines",
    "skill_description": "I am a passionate Web Developer...",
    "experiences": [
        {
            "company_name": "Acme Inc.",
            "location": "Urban Rodriguez, Rizal",
            "job_title": "Full Stack Web Developer",
            "date_from": "October 2025",
            "date_to": "Present",
            "bullet_points": [
                "Experienced with Java and Spring Boot...",
                "Utilized Docker and Kubernetes..."
            ]
        }
    ],
    "educations": [
        {
            "school": "Colegio De Montalban",
            "location": "Rodriguez, Rizal, Philippines",
            "school_type": "B.S in Information Technology",
            "date_from": "2021",
            "date_to": "2025"
        }
    ],
    "skills": ["ROI Modeling", "Cost-Benefit Analysis", "Inventory Management"]
}
